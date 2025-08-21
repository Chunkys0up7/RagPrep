"""
Document Parsers

This module implements the document parsing system with a cascading parser strategy.
It provides parsers for various document formats with fallback mechanisms and error handling.
"""

from abc import ABC, abstractmethod
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple, Union

import logging
import os
import traceback
from pathlib import Path

from config import Config, get_config

logger = logging.getLogger(__name__)


@dataclass
class ParsedContent:
    """Container for parsed document content."""

    text_content: str
    structured_content: Dict[str, Any]
    metadata: Dict[str, Any]
    tables: List[Dict[str, Any]]
    images: List[Dict[str, Any]]
    math_content: List[Dict[str, Any]]
    structure: Dict[str, Any]
    parser_used: str
    parsing_errors: List[str]
    parsing_warnings: List[str]


@dataclass
class ParserResult:
    """Result of document parsing operation."""

    success: bool
    content: Optional[ParsedContent]
    error_message: Optional[str]
    parser_name: str
    processing_time: float
    metadata: Dict[str, Any]


class DocumentParser(ABC):
    """Abstract base class for document parsers."""

    def __init__(self, config: Config):
        """Initialize parser with configuration."""
        self.config = config
        self.supported_formats: List[str] = []
        self.parser_name: str = self.__class__.__name__

    @abstractmethod
    def can_parse(self, file_path: str) -> bool:
        """Check if this parser can handle the given file."""
        pass

    @abstractmethod
    def parse(self, file_path: str) -> ParserResult:
        """Parse the document and return structured content."""
        pass

    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract basic metadata from the file."""
        path = Path(file_path)
        return {
            "filename": path.name,
            "file_extension": path.suffix.lower(),
            "file_size": path.stat().st_size if path.exists() else 0,
            "parser": self.parser_name,
        }


class PDFParser(DocumentParser):
    """PDF document parser using PyMuPDF as primary parser."""

    def __init__(self, config: Config):
        """Initialize PDF parser."""
        super().__init__(config)
        self.supported_formats = [".pdf"]
        self.parser_name = "PDFParser"

        # Initialize PyMuPDF
        try:
            import fitz  # PyMuPDF

            self.fitz_available = True
        except ImportError:
            logger.warning("PyMuPDF not available, PDF parsing will be limited")
            self.fitz_available = False

    def can_parse(self, file_path: str) -> bool:
        """Check if this parser can handle the PDF file."""
        return Path(file_path).suffix.lower() == ".pdf" and self.fitz_available

    def parse(self, file_path: str) -> ParserResult:
        """Parse PDF document using PyMuPDF."""
        import time

        start_time = time.time()

        try:
            if not self.can_parse(file_path):
                return ParserResult(
                    success=False,
                    content=None,
                    error_message="Cannot parse this file type",
                    parser_name=self.parser_name,
                    processing_time=time.time() - start_time,
                    metadata={},
                )

            import fitz  # PyMuPDF

            # Open the PDF
            doc = fitz.open(file_path)

            # Extract text content
            text_content = ""
            structured_content = {
                "pages": [],
                "tables": [],
                "images": [],
                "annotations": [],
            }

            # Process each page
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)

                # Extract text
                page_text = page.get_text()
                text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"

                # Extract page structure
                page_structure = {
                    "page_number": page_num + 1,
                    "text_blocks": page.get_text("dict")["blocks"],
                    "images": page.get_images(),
                    "annotations": page.annots(),
                }

                structured_content["pages"].append(page_structure)

                # Extract tables (basic extraction)
                tables = self._extract_tables_from_page(page)
                structured_content["tables"].extend(tables)

                # Extract images
                images = self._extract_images_from_page(page)
                structured_content["images"].extend(images)

            # Extract document metadata
            metadata = doc.metadata
            metadata.update(self.get_metadata(file_path))

            # Create parsed content
            content = ParsedContent(
                text_content=text_content.strip(),
                structured_content=structured_content,
                metadata=metadata,
                tables=structured_content["tables"],
                images=structured_content["images"],
                math_content=[],  # PDF math extraction would require additional processing
                structure={"pages": len(doc), "total_pages": len(doc)},
                parser_used=self.parser_name,
                parsing_errors=[],
                parsing_warnings=[],
            )

            doc.close()

            return ParserResult(
                success=True,
                content=content,
                error_message=None,
                parser_name=self.parser_name,
                processing_time=time.time() - start_time,
                metadata=metadata,
            )

        except Exception as e:
            error_msg = f"Error parsing PDF {file_path}: {str(e)}"
            logger.error(error_msg)
            logger.debug(traceback.format_exc())

            return ParserResult(
                success=False,
                content=None,
                error_message=error_msg,
                parser_name=self.parser_name,
                processing_time=time.time() - start_time,
                metadata=self.get_metadata(file_path),
            )

    def _extract_tables_from_page(self, page) -> List[Dict[str, Any]]:
        """Extract tables from a PDF page."""
        tables = []
        try:
            # Basic table extraction - this is a simplified approach
            # In production, you might want to use more sophisticated table detection
            text_dict = page.get_text("dict")

            for block in text_dict["blocks"]:
                if "lines" in block:
                    # Simple heuristic: if block has multiple lines with similar structure
                    if len(block["lines"]) > 1:
                        table_data = []
                        for line in block["lines"]:
                            row = []
                            for span in line["spans"]:
                                row.append(span["text"])
                            if row:
                                table_data.append(row)

                        if (
                            len(table_data) > 1
                        ):  # At least 2 rows to be considered a table
                            tables.append(
                                {
                                    "data": table_data,
                                    "bbox": block["bbox"],
                                    "confidence": 0.7,  # Basic confidence
                                }
                            )
        except Exception as e:
            logger.debug(f"Error extracting tables from page: {e}")

        return tables

    def _extract_images_from_page(self, page) -> List[Dict[str, Any]]:
        """Extract images from a PDF page."""
        images = []
        try:
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    pix = page.parent.extract_image(xref)
                    if pix:
                        images.append(
                            {
                                "index": img_index,
                                "bbox": img[1:5],
                                "width": img[2] - img[0],
                                "height": img[3] - img[1],
                                "format": pix["ext"],
                                "size": len(pix["image"]),
                            }
                        )
                except Exception as e:
                    logger.debug(f"Error extracting image {img_index}: {e}")
        except Exception as e:
            logger.debug(f"Error extracting images from page: {e}")

        return images


class DOCXParser(DocumentParser):
    """DOCX document parser using python-docx."""

    def __init__(self, config: Config):
        """Initialize DOCX parser."""
        super().__init__(config)
        self.supported_formats = [".docx", ".doc"]
        self.parser_name = "DOCXParser"

        try:
            from docx import Document

            self.docx_available = True
        except ImportError:
            logger.warning("python-docx not available, DOCX parsing will be limited")
            self.docx_available = False

    def can_parse(self, file_path: str) -> bool:
        """Check if this parser can handle the DOCX file."""
        return (
            Path(file_path).suffix.lower() in [".docx", ".doc"] and self.docx_available
        )

    def parse(self, file_path: str) -> ParserResult:
        """Parse DOCX document."""
        import time

        start_time = time.time()

        try:
            if not self.can_parse(file_path):
                return ParserResult(
                    success=False,
                    content=None,
                    error_message="Cannot parse this file type",
                    parser_name=self.parser_name,
                    processing_time=time.time() - start_time,
                    metadata={},
                )

            from docx import Document

            # Open the document
            doc = Document(file_path)

            # Extract text content
            text_content = ""
            structured_content = {
                "paragraphs": [],
                "tables": [],
                "sections": [],
                "headers": [],
                "footers": [],
            }

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content += para.text + "\n"
                    structured_content["paragraphs"].append(
                        {
                            "text": para.text,
                            "style": para.style.name,
                            "runs": [run.text for run in para.runs],
                        }
                    )

            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)

                structured_content["tables"].append(
                    {
                        "data": table_data,
                        "rows": len(table_data),
                        "columns": len(table_data[0]) if table_data else 0,
                    }
                )

            # Extract document properties
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title,
                "author": core_props.author,
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
                "revision": core_props.revision,
                "category": core_props.category,
                "comments": core_props.comments,
            }
            metadata.update(self.get_metadata(file_path))

            # Create parsed content
            content = ParsedContent(
                text_content=text_content.strip(),
                structured_content=structured_content,
                metadata=metadata,
                tables=structured_content["tables"],
                images=[],  # DOCX image extraction would require additional processing
                math_content=[],  # DOCX math extraction would require additional processing
                structure={
                    "paragraphs": len(structured_content["paragraphs"]),
                    "tables": len(structured_content["tables"]),
                },
                parser_used=self.parser_name,
                parsing_errors=[],
                parsing_warnings=[],
            )

            return ParserResult(
                success=True,
                content=content,
                error_message=None,
                parser_name=self.parser_name,
                processing_time=time.time() - start_time,
                metadata=metadata,
            )

        except Exception as e:
            error_msg = f"Error parsing DOCX {file_path}: {str(e)}"
            logger.error(error_msg)
            logger.debug(traceback.format_exc())

            return ParserResult(
                success=False,
                content=None,
                error_message=error_msg,
                parser_name=self.parser_name,
                processing_time=time.time() - start_time,
                metadata=self.get_metadata(file_path),
            )


class TextParser(DocumentParser):
    """Simple text file parser."""

    def __init__(self, config: Config):
        """Initialize text parser."""
        super().__init__(config)
        self.supported_formats = [".txt", ".md", ".csv", ".log"]
        self.parser_name = "TextParser"

    def can_parse(self, file_path: str) -> bool:
        """Check if this parser can handle the text file."""
        return Path(file_path).suffix.lower() in self.supported_formats

    def parse(self, file_path: str) -> ParserResult:
        """Parse text document."""
        import time

        start_time = time.time()

        try:
            if not self.can_parse(file_path):
                return ParserResult(
                    success=False,
                    content=None,
                    error_message="Cannot parse this file type",
                    parser_name=self.parser_name,
                    processing_time=time.time() - start_time,
                    metadata={},
                )

            # Read text content
            with open(file_path, "r", encoding="utf-8") as f:
                text_content = f.read()

            # Basic structure analysis
            lines = text_content.split("\n")
            paragraphs = [line.strip() for line in lines if line.strip()]

            structured_content = {
                "lines": lines,
                "paragraphs": paragraphs,
                "line_count": len(lines),
                "word_count": len(text_content.split()),
                "character_count": len(text_content),
            }

            # Extract metadata
            metadata = self.get_metadata(file_path)
            metadata.update(
                {
                    "encoding": "utf-8",
                    "line_count": len(lines),
                    "word_count": len(text_content.split()),
                    "character_count": len(text_content),
                }
            )

            # Create parsed content
            content = ParsedContent(
                text_content=text_content,
                structured_content=structured_content,
                metadata=metadata,
                tables=[],
                images=[],
                math_content=[],
                structure={"lines": len(lines), "paragraphs": len(paragraphs)},
                parser_used=self.parser_name,
                parsing_errors=[],
                parsing_warnings=[],
            )

            return ParserResult(
                success=True,
                content=content,
                error_message=None,
                parser_name=self.parser_name,
                processing_time=time.time() - start_time,
                metadata=metadata,
            )

        except Exception as e:
            error_msg = f"Error parsing text file {file_path}: {str(e)}"
            logger.error(error_msg)
            logger.debug(traceback.format_exc())

            return ParserResult(
                success=False,
                content=None,
                error_message=error_msg,
                parser_name=self.parser_name,
                processing_time=time.time() - start_time,
                metadata=self.get_metadata(file_path),
            )


class HTMLParser(DocumentParser):
    """HTML document parser using BeautifulSoup."""

    def __init__(self, config: Config):
        """Initialize HTML parser."""
        super().__init__(config)
        self.supported_formats = [".html", ".htm"]
        self.parser_name = "HTMLParser"

        try:
            from bs4 import BeautifulSoup

            self.bs4_available = True
        except ImportError:
            logger.warning("BeautifulSoup not available, HTML parsing will be limited")
            self.bs4_available = False

    def can_parse(self, file_path: str) -> bool:
        """Check if this parser can handle the HTML file."""
        return (
            Path(file_path).suffix.lower() in [".html", ".htm"] and self.bs4_available
        )

    def parse(self, file_path: str) -> ParserResult:
        """Parse HTML document."""
        import time

        start_time = time.time()

        try:
            if not self.can_parse(file_path):
                return ParserResult(
                    success=False,
                    content=None,
                    error_message="Cannot parse this file type",
                    parser_name=self.parser_name,
                    processing_time=time.time() - start_time,
                    metadata={},
                )

            from bs4 import BeautifulSoup

            # Read HTML content
            with open(file_path, "r", encoding="utf-8") as f:
                html_content = f.read()

            # Parse with BeautifulSoup
            soup = BeautifulSoup(html_content, "html.parser")

            # Extract text content
            text_content = soup.get_text(separator="\n", strip=True)

            # Extract structured content
            structured_content = {
                "title": soup.title.string if soup.title else None,
                "headings": [],
                "paragraphs": [],
                "links": [],
                "tables": [],
                "images": [],
                "lists": [],
            }

            # Extract headings
            for tag in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                for heading in soup.find_all(tag):
                    structured_content["headings"].append(
                        {
                            "level": int(tag[1]),
                            "text": heading.get_text(strip=True),
                            "id": heading.get("id"),
                        }
                    )

            # Extract paragraphs
            for p in soup.find_all("p"):
                structured_content["paragraphs"].append(p.get_text(strip=True))

            # Extract links
            for a in soup.find_all("a", href=True):
                structured_content["links"].append(
                    {"text": a.get_text(strip=True), "href": a["href"]}
                )

            # Extract tables
            for table in soup.find_all("table"):
                table_data = []
                for row in table.find_all("tr"):
                    row_data = []
                    for cell in row.find_all(["td", "th"]):
                        row_data.append(cell.get_text(strip=True))
                    if row_data:
                        table_data.append(row_data)

                if table_data:
                    structured_content["tables"].append(
                        {
                            "data": table_data,
                            "rows": len(table_data),
                            "columns": len(table_data[0]) if table_data else 0,
                        }
                    )

            # Extract images
            for img in soup.find_all("img"):
                structured_content["images"].append(
                    {
                        "src": img.get("src"),
                        "alt": img.get("alt"),
                        "title": img.get("title"),
                    }
                )

            # Extract lists
            for ul in soup.find_all(["ul", "ol"]):
                list_items = [li.get_text(strip=True) for li in ul.find_all("li")]
                structured_content["lists"].append(
                    {"type": ul.name, "items": list_items}
                )

            # Extract metadata
            metadata = self.get_metadata(file_path)
            metadata.update(
                {
                    "title": structured_content["title"],
                    "headings_count": len(structured_content["headings"]),
                    "paragraphs_count": len(structured_content["paragraphs"]),
                    "links_count": len(structured_content["links"]),
                    "tables_count": len(structured_content["tables"]),
                    "images_count": len(structured_content["images"]),
                }
            )

            # Create parsed content
            content = ParsedContent(
                text_content=text_content,
                structured_content=structured_content,
                metadata=metadata,
                tables=structured_content["tables"],
                images=structured_content["images"],
                math_content=[],
                structure={
                    "headings": len(structured_content["headings"]),
                    "paragraphs": len(structured_content["paragraphs"]),
                    "tables": len(structured_content["tables"]),
                },
                parser_used=self.parser_name,
                parsing_errors=[],
                parsing_warnings=[],
            )

            return ParserResult(
                success=True,
                content=content,
                error_message=None,
                parser_name=self.parser_name,
                processing_time=time.time() - start_time,
                metadata=metadata,
            )

        except Exception as e:
            error_msg = f"Error parsing HTML {file_path}: {str(e)}"
            logger.error(error_msg)
            logger.debug(traceback.format_exc())

            return ParserResult(
                success=False,
                content=None,
                error_message=error_msg,
                parser_name=self.parser_name,
                processing_time=time.time() - start_time,
                metadata=self.get_metadata(file_path),
            )


class CascadingDocumentParser:
    """Main document parser that uses a cascading strategy with multiple parsers."""

    def __init__(self, config: Config):
        """Initialize the cascading parser with configuration."""
        self.config = config
        self.parsers: List[DocumentParser] = []
        self._initialize_parsers()

    def _initialize_parsers(self):
        """Initialize available parsers based on configuration."""
        # Initialize parsers in order of preference
        self.parsers = [
            PDFParser(self.config),
            DOCXParser(self.config),
            HTMLParser(self.config),
            TextParser(self.config),
        ]

        # Filter out unavailable parsers
        self.parsers = [
            parser for parser in self.parsers if self._is_parser_available(parser)
        ]

        logger.info(f"Initialized {len(self.parsers)} document parsers")

    def _is_parser_available(self, parser: DocumentParser) -> bool:
        """Check if a parser is available and functional."""
        try:
            # Try to create a minimal test to check parser availability
            if hasattr(parser, "fitz_available") and not parser.fitz_available:
                return False
            if hasattr(parser, "docx_available") and not parser.docx_available:
                return False
            if hasattr(parser, "bs4_available") and not parser.bs4_available:
                return False
            return True
        except Exception:
            return False

    def get_available_parsers(self) -> List[str]:
        """Get list of available parser names."""
        return [parser.parser_name for parser in self.parsers]

    def get_supported_formats(self) -> List[str]:
        """Get list of all supported file formats."""
        formats = set()
        for parser in self.parsers:
            formats.update(parser.supported_formats)
        return sorted(list(formats))

    def parse(self, file_path: str) -> ParserResult:
        """Parse document using cascading parser strategy (alias for parse_document)."""
        return self.parse_document(file_path)

    def parse_document(self, file_path: str) -> ParserResult:
        """Parse document using cascading parser strategy."""
        file_path = str(file_path)

        if not os.path.exists(file_path):
            return ParserResult(
                success=False,
                content=None,
                error_message=f"File not found: {file_path}",
                parser_name="CascadingParser",
                processing_time=0.0,
                metadata={},
            )

        # Get file extension
        file_ext = Path(file_path).suffix.lower()

        # Find parsers that can handle this file type
        suitable_parsers = [
            parser for parser in self.parsers if parser.can_parse(file_path)
        ]

        if not suitable_parsers:
            return ParserResult(
                success=False,
                content=None,
                error_message=f"No suitable parser found for file type: {file_ext}",
                parser_name="CascadingParser",
                processing_time=0.0,
                metadata={"file_extension": file_ext},
            )

        # Try parsers in order until one succeeds
        for parser in suitable_parsers:
            logger.info(f"Attempting to parse {file_path} with {parser.parser_name}")

            try:
                result = parser.parse(file_path)

                if result.success:
                    logger.info(
                        f"Successfully parsed {file_path} with {parser.parser_name}"
                    )
                    return result
                else:
                    logger.warning(
                        f"Parser {parser.parser_name} failed: {result.error_message}"
                    )

            except Exception as e:
                logger.error(f"Error with parser {parser.parser_name}: {str(e)}")
                continue

        # If all parsers failed, return error result
        return ParserResult(
            success=False,
            content=None,
            error_message=f"All parsers failed to parse {file_path}",
            parser_name="CascadingParser",
            processing_time=0.0,
            metadata={
                "file_extension": file_ext,
                "attempted_parsers": [p.parser_name for p in suitable_parsers],
            },
        )

    def parse_batch(self, file_paths: List[str]) -> List[ParserResult]:
        """Parse multiple documents in batch."""
        results = []

        for file_path in file_paths:
            logger.info(f"Processing batch file: {file_path}")
            result = self.parse_document(file_path)
            results.append(result)

        return results


# Convenience function to get parser instance
def get_document_parser(config: Optional[Config] = None) -> CascadingDocumentParser:
    """Get a configured document parser instance."""
    if config is None:
        config = get_config()
    return CascadingDocumentParser(config)
